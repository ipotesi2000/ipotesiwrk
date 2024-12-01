#!/usr/bin/env python
# coding: utf-8
import os
import pathlib as pt
from pathlib import Path
import sys
import markdownify
from docx import Document
import shutil

# pip install python-docx markdownify

def rlist(path: str, match: str = "*") -> list:
    return [str(x) for x in pt.Path(path).rglob(match)]

def make_dir(path: str, mode: int = 0o777):
    p = pt.Path(path)
    p.mkdir(parents=True, exist_ok=True, mode=mode)
    p.chmod(mode=mode)


def move2back(file_path):
    # Converti il percorso del file in un oggetto Path
    file_path = Path(file_path)
    # Ottieni la directory corrente del file
    current_dir = file_path.parent
    # Crea il nome della directory di backup
    backup_dir = current_dir.with_name(current_dir.name + "_back")
    # Crea la directory di backup se non esiste
    backup_dir.mkdir(parents=True, exist_ok=True)
    # Costruisci il percorso di destinazione nella directory di backup
    destination_path = backup_dir / file_path.name
    # Sposta il file nella directory di backup
    shutil.move(str(file_path), str(destination_path))

##########################

def docx2md(docx_path, md_path):
    try:
        doc = Document(docx_path)
        md_content = ""

        def convert_paragraph(para):
            md_text = ""
            for run in para.runs:
                text = run.text
                if run.bold:
                    text = f"**{text}**"
                if run.italic:
                    text = f"*{text}*"
                if run.underline:
                    text = f"<u>{text}</u>"
                if run.element.xpath('.//w:hyperlink'):
                    hyperlink = run.element.xpath('.//w:hyperlink')[0]
                    url = hyperlink.get(qn('w:val'))
                    text = f"[{text}]({url})"
                md_text += text
            return md_text

        def convert_table(table):
            md_table = ""
            for row in table.rows:
                md_table += "| " + " | ".join([convert_paragraph(cell) for cell in row.cells]) + " |\n"
            return md_table

        def convert_list(para):
            md_text = ""
            if para.style.name.startswith('List Paragraph'):
                if para.style.name.endswith('Bullet'):
                    md_text += "- "
                elif para.style.name.endswith('Number'):
                    md_text += f"{para.style.level + 1}. "
                md_text += convert_paragraph(para)
            return md_text

        def convert_image(para):
            md_text = ""
            for run in para.runs:
                if run.element.xpath('.//w:drawing'):
                    blip = run.element.xpath('.//a:blip', namespaces=run.element.nsmap)[0]
                    rId = blip.get(qn('r:embed'))
                    image_part = doc.part.related_parts[rId]
                    image_path = os.path.join(os.path.dirname(docx_path), image_part.partname.split('/')[-1])
                    with open(image_path, 'wb') as f:
                        f.write(image_part.blob)
                    md_text += f"![Image]({image_path})\n\n"
            return md_text

        for para in doc.paragraphs:
            if para.style.name.startswith('List Paragraph'):
                md_content += convert_list(para) + "\n"
            elif para.text.strip() == "":
                md_content += "\n"
            else:
                md_content += convert_paragraph(para) + "\n\n"

        for table in doc.tables:
            md_content += "\n" + convert_table(table) + "\n"

        for para in doc.paragraphs:
            md_content += convert_image(para)

        with open(md_path, 'w', encoding='utf-8') as md_file:
            md_file.write(md_content)

    except Exception as e:
        print(f"Errore: {e}")
        exit(1)

def dirdocx2md(src_dir, dest_dir):
    make_dir(dest_dir)
    path_lst = rlist(src_dir, "*.docx")
    for doc_path in path_lst:
        print(doc_path)
        fname = os.path.basename(doc_path).replace(".docx", ".md")
        fname = fname.lower()
        fname = fname.replace(" ", "_")
        md_path = os.path.join(dest_dir, fname)
        docx2md(doc_path, md_path)
        move2back(doc_path)



if __name__ == "__main__":
    dir_src = "./data/docx"
    dir_dst = "./data/md"
    if len(sys.argv) < 3:
     print("\n\nread_prd.py <dir sorgente>  <dir destinazione>")
     exit()
    dir_src = sys.argv[1]
    dir_dst = sys.argv[2]
    dirdocx2md(dir_src, dir_dst)
